import base64
import json
import os
from typing import List, Dict, Any, Optional
from openai import OpenAI
from PIL import Image
import io
import fitz  # PyMuPDF
from docx import Document
import pandas as pd

class AIProcessor:
    def __init__(self, api_key: str, base_url: str = "https://api.aitunnel.ru/v1/"):
        self.client = OpenAI(api_key=api_key, base_url=base_url)
        self.model = "gemini-3-flash-preview"
        self.system_instruction = (
            "Ты — строгий ассистент по структурированию данных. Твоя задача — извлекать информацию о компаниях "
            "из предоставленных материалов и возвращать её СТРОГО в формате JSON-списка.\n\n"
            "КРИТИЧЕСКИЕ ПРАВИЛА:\n"
            "1. ЗАПРЕЩЕНО ВЫДУМЫВАТЬ ИНФОРМАЦИЮ. Если в тексте или на изображении нет конкретных данных (например, телефона, ИНН или сайта), "
            "в соответствующем поле JSON ОБЯЗАТЕЛЬНО ставь null.\n"
            "2. Не используй общие знания о компаниях. Опирайся ТОЛЬКО на предоставленный контент.\n"
            "3. Если информация сомнительна или неполна, лучше поставить null, чем гадать.\n"
            "4. В поле 'Заметки' ЗАПРЕЩЕНО писать выводы, резюме или описания от себя. Это поле должно оставаться null, если в исходном документе нет специфических пометок, не вошедших в другие поля.\n\n"
            "Поля JSON:\n"
            "- Наименование предприятия: Полное название из документа.\n"
            "- ИНН: Только цифры (если есть).\n"
            "- Юр. форма: ООО, ИП, АНО и т.д.\n"
            "- Изделия: Перечисли через ';' товары или услуги, упомянутые в тексте.\n"
            "- Группа контрагентов: Если упомянута.\n"
            "- ФИО руководителя: Полное имя руководителя.\n"
            "- Должность руководителя: Например, Генеральный директор.\n"
            "- Адрес: Полный адрес.\n"
            "- Часовой пояс: Определи часовой пояс на основе адреса. Формат: '+3', '+5', '-2' и т.д. (относительно UTC).\n"
            "- Телефоны: Формат +7 (XXX) XXX-XX-XX. Если несколько — через ';'. ТОЛЬКО ИЗ ТЕКСТА.\n"
            "- Сайт: URL сайта.\n"
            "- Почта для рассылки: Email.\n"
            "- Рубрика: Определи категорию бизнеса самостоятельно на основе описания.\n"
            "- Подрубрика: Более узкая категория.\n"
            "- Заметки: null (НЕ ЗАПОЛНЯТЬ, если нет явных пометок в тексте)."
        )

    def encode_image(self, image_path: str) -> str:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        text = ""
        try:
            doc = fitz.open(pdf_path)
            for page in doc:
                text += page.get_text()
            doc.close()
        except Exception as e:
            text = f"[Ошибка чтения PDF {pdf_path}: {e}]"
        return text

    def extract_text_from_docx(self, docx_path: str) -> str:
        try:
            doc = Document(docx_path)
            full_text = []
            
            # Извлекаем текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                full_text.append("\n--- Таблица ---")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        full_text.append(" | ".join(row_data))
                full_text.append("---------------\n")
                
            return "\n".join(full_text)
        except Exception as e:
            return f"[Ошибка чтения DOCX {docx_path}: {e}]"

    def extract_text_from_excel(self, xlsx_path: str) -> str:
        try:
            df = pd.read_excel(xlsx_path)
            return df.to_string()
        except Exception as e:
            return f"[Ошибка чтения Excel {xlsx_path}: {e}]"

    def process_content(self, text: Optional[str] = None, file_paths: Optional[List[str]] = None) -> List[Dict[str, Any]]:
        messages = [
            {"role": "system", "content": self.system_instruction}
        ]
        
        user_content = []
        combined_text = text if text else ""
        
        if file_paths:
            for path in file_paths:
                ext = os.path.splitext(path)[1].lower()
                
                if ext in [".jpg", ".jpeg", ".png", ".webp"]:
                    base64_image = self.encode_image(path)
                    mime_type = "image/jpeg" if ext in [".jpg", ".jpeg"] else f"image/{ext[1:]}"
                    user_content.append({
                        "type": "image_url",
                        "image_url": {"url": f"data:{mime_type};base64,{base64_image}"}
                    })
                elif ext == ".pdf":
                    combined_text += f"\n--- Содержимое PDF ({os.path.basename(path)}) ---\n"
                    combined_text += self.extract_text_from_pdf(path)
                elif ext == ".docx":
                    combined_text += f"\n--- Содержимое Word ({os.path.basename(path)}) ---\n"
                    combined_text += self.extract_text_from_docx(path)
                elif ext in [".xlsx", ".xls"]:
                    combined_text += f"\n--- Содержимое Excel ({os.path.basename(path)}) ---\n"
                    combined_text += self.extract_text_from_excel(path)
                elif ext == ".txt":
                    try:
                        with open(path, 'r', encoding='utf-8') as f:
                            combined_text += f"\n--- Содержимое TXT ({os.path.basename(path)}) ---\n"
                            combined_text += f.read()
                    except:
                        with open(path, 'r', encoding='cp1251') as f:
                            combined_text += f.read()

        if combined_text.strip():
            user_content.append({"type": "text", "text": f"Обработай следующие данные и извлеки информацию о компаниях:\n\n{combined_text}"})
        
        if not user_content:
            return []

        messages.append({"role": "user", "content": user_content})

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=messages,
                response_format={"type": "json_object"}
            )
            
            content = response.choices[0].message.content
            data = json.loads(content)
            
            if isinstance(data, dict):
                for key in data:
                    if isinstance(data[key], list):
                        return data[key]
                return [data]
            return data if isinstance(data, list) else []
            
        except Exception as e:
            print(f"Error processing with AI: {e}")
            raise e
